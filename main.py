import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from ttkbootstrap import Style
import datetime
import os
import sys
import json
import csv
import matplotlib.pyplot as plt
from collections import Counter
from openpyxl import Workbook
from docx import Document
from docx2pdf import convert
import docx.shared

from config import MUSHROOM_TYPES, RESTAURANT_ASSIGNMENTS

LOG_FILE = "logs.json"

class MushroomApp:
    def __init__(self, root):
        self.root = root
        self.root.title("🍄 Mushroom Tracking System")
        self.current_theme = "darkly"
        self.style = Style(self.current_theme)

        self.logs = []
        self.filtered_logs = []
        self.is_mock_mode = os.getenv("USE_MOCK_SQUARE", "1") == "1"
        self.settings = {
            "theme": "darkly",
            "default_restaurant_id": 1,
            "invoice_template": ""
        }
        self.settings_file = "settings.json"
        self.load_settings()

        self.build_gui()
        self.load_logs()
        self.add_theme_toggle_button()
        
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

    def on_close(self):
        self.save_logs()
        self.root.destroy()

    def add_theme_toggle_button(self):
        toggle_frame = ttk.Frame(self.root)
        toggle_frame.pack(pady=(5, 10))
        ttk.Button(toggle_frame, text="Toggle Theme", command=self.toggle_theme).pack()

    def toggle_theme(self):
        self.current_theme = "flatly" if self.current_theme == "darkly" else "darkly"
        self.style.theme_use(self.current_theme)
        self.settings["theme"] = self.current_theme  # Save new theme
        self.save_settings()
        self.show_toast(f"Switched to {self.current_theme.capitalize()} Mode!", "info")

    def build_gui(self):
        main_frame = ttk.Frame(self.root, padding=20)
        main_frame.pack(expand=True, fill="both")

        # --- Search Field ---
        search_frame = ttk.Frame(main_frame)
        search_frame.pack(fill="x", pady=(0, 10))

        ttk.Label(search_frame, text="Search Deliveries:").pack(side="left")
        self.search_var = tk.StringVar()
        self.search_var.trace_add("write", self.update_filtered_logs)
        ttk.Entry(search_frame, textvariable=self.search_var, width=40).pack(side="left", padx=10)

        # --- Date Range Filter ---
        date_filter_frame = ttk.Frame(main_frame)
        date_filter_frame.pack(fill="x", pady=(5, 10))

        ttk.Label(date_filter_frame, text="Start Date (YYYY-MM-DD):").pack(side="left")
        self.start_date_var = tk.StringVar()
        self.start_date_var.trace_add("write", self.update_filtered_logs)
        ttk.Entry(date_filter_frame, textvariable=self.start_date_var, width=12).pack(side="left", padx=5)

        ttk.Label(date_filter_frame, text="End Date (YYYY-MM-DD):").pack(side="left", padx=(20, 0))
        self.end_date_var = tk.StringVar()
        self.end_date_var.trace_add("write", self.update_filtered_logs)
        ttk.Entry(date_filter_frame, textvariable=self.end_date_var, width=12).pack(side="left", padx=5)

        # --- Delivery Form ---
        form_frame = ttk.LabelFrame(main_frame, text="Delivery Information", padding=(20, 10))
        form_frame.pack(fill="x")

        self.mushroom_type_var = tk.StringVar()
        ttk.Label(form_frame, text="Mushroom Type:").grid(row=0, column=0, sticky="w", pady=5)
        self.mushroom_dropdown = ttk.Combobox(
            form_frame, textvariable=self.mushroom_type_var,
            values=[f"{k} - {v}" for k, v in MUSHROOM_TYPES.items()], state="readonly"
        )
        self.mushroom_dropdown.grid(row=0, column=1, padx=10, pady=5, sticky="ew")

        self.box_number_var = tk.StringVar()
        ttk.Label(form_frame, text="Box Number:").grid(row=1, column=0, sticky="w", pady=5)
        ttk.Entry(form_frame, textvariable=self.box_number_var).grid(row=1, column=1, padx=10, pady=5, sticky="ew")

        self.restaurant_id_var = tk.StringVar()
        ttk.Label(form_frame, text="Restaurant ID:").grid(row=2, column=0, sticky="w", pady=5)
        self.restaurant_dropdown = ttk.Combobox(
            form_frame, textvariable=self.restaurant_id_var,
            values=[f"{k} - {v}" for k, v in RESTAURANT_ASSIGNMENTS.items()], state="readonly"
        )
        self.restaurant_dropdown.grid(row=2, column=1, padx=10, pady=5, sticky="ew")

        # Pre-select default restaurant if settings available
        default_id = self.settings.get("default_restaurant_id")
        if default_id:
            for i, value in enumerate(self.restaurant_dropdown['values']):
                if value.startswith(f"{default_id} -"):
                    self.restaurant_dropdown.current(i)
                    break

        self.pack_date_var = tk.StringVar()
        ttk.Label(form_frame, text="Pack Date (YYYY-MM-DD):").grid(row=3, column=0, sticky="w", pady=5)
        self.pack_date_entry = ttk.Entry(form_frame, textvariable=self.pack_date_var)
        self.pack_date_entry.grid(row=3, column=1, padx=10, pady=5, sticky="ew")
        ttk.Button(form_frame, text="Pick", command=lambda: self.pick_date(self.pack_date_var)).grid(row=3, column=2, padx=5)

        self.ship_date_var = tk.StringVar()
        ttk.Label(form_frame, text="Ship Date (YYYY-MM-DD):").grid(row=4, column=0, sticky="w", pady=5)
        self.ship_date_entry = ttk.Entry(form_frame, textvariable=self.ship_date_var)
        self.ship_date_entry.grid(row=4, column=1, padx=10, pady=5, sticky="ew")
        ttk.Button(form_frame, text="Pick", command=lambda: self.pick_date(self.ship_date_var)).grid(row=4, column=2, padx=5)

        # --- Action Buttons ---
        button_frame = ttk.LabelFrame(main_frame, text="Actions", padding=(20, 10))
        button_frame.pack(fill="x", pady=(15, 0))

        actions = [
            ("Add Entry", self.confirm_add),
            ("Clear Form", self.clear_form),
            ("View Log", self.view_log),
            ("Save Logs", self.save_logs),
            ("Load Logs", self.load_logs),
            ("Generate Invoice (PDF)", self.generate_invoice),
            ("Show Charts", self.show_charts),
            ("Toggle Mock/Live Mode", self.toggle_mode),
            ("Settings", self.open_settings_window),
            ("Clear All Logs", self.clear_logs),
            ("Backup Manager", self.open_backup_manager),
            ("Export Summary Report", self.export_summary_report),
            ("Edit Logs", self.edit_logs),
        ]

        # Place Action Buttons
        for i, (text, command) in enumerate(actions):
            ttk.Button(button_frame, text=text, command=command).grid(row=i//3, column=i%3, padx=10, pady=10, sticky="ew")

        # Special Export Button — needs self reference for enabling/disabling
        self.export_button = ttk.Button(button_frame, text="Export Data", command=self.export_data)
        self.export_button.grid(row=len(actions)//3, column=len(actions)%3, padx=10, pady=10, sticky="ew")

        # --- Status Display ---
        status_frame = ttk.Frame(main_frame, padding=(0, 10))
        status_frame.pack(fill="x")
        self.status_label = ttk.Label(status_frame, text=self.get_mode_text(), font=("Segoe UI", 10, "italic"))
        self.status_label.pack()

        # Stretch layout
        for frame in [form_frame, button_frame]:
            frame.columnconfigure(1, weight=1)

    def pick_date(self, target_var):
        top = tk.Toplevel(self.root)
        top.title("Pick a Date")

        today = datetime.date.today()
        dates = [(today + datetime.timedelta(days=i)).strftime("%Y-%m-%d") for i in range(0, 365)]

        listbox = tk.Listbox(top, width=20, height=10)
        listbox.pack(padx=10, pady=10)

        for date in dates:
            listbox.insert(tk.END, date)

        def set_selected_date():
            try:
                selected = listbox.get(listbox.curselection())
                target_var.set(selected)
                top.destroy()
            except:
                messagebox.showwarning("Selection Error", "Please select a date.")

        ttk.Button(top, text="Select", command=set_selected_date).pack(pady=5)

    def confirm_add(self):
        if not self.validate_inputs():
            return
        if messagebox.askyesno("Confirm", "Add this entry to the traceability log?"):
            label = self.generate_label()
            self.logs.append(label)
            self.save_logs()
            self.update_filtered_logs()
            self.update_export_button_state()  # 🔥 here
            self.show_toast(f"Added:\n{label}", "success")
            self.clear_form()

    def validate_inputs(self):
        try:
            int(self.mushroom_type_var.get().split(" - ")[0])
            int(self.box_number_var.get())
            int(self.restaurant_id_var.get().split(" - ")[0])
            datetime.datetime.strptime(self.pack_date_var.get(), "%Y-%m-%d")
            datetime.datetime.strptime(self.ship_date_var.get(), "%Y-%m-%d")
            return True
        except Exception as e:
            messagebox.showerror("Input Error", f"Invalid input: {e}")
            return False

    def generate_label(self):
        mushroom_id = int(self.mushroom_type_var.get().split(" - ")[0])
        box_number = int(self.box_number_var.get())
        restaurant_id = int(self.restaurant_id_var.get().split(" - ")[0])
        pack_date = self.pack_date_var.get()
        ship_date = self.ship_date_var.get()
        mushroom_name = MUSHROOM_TYPES[mushroom_id]
        restaurant_name = RESTAURANT_ASSIGNMENTS[restaurant_id]
        tracking_number = f"{pack_date}-BOX{box_number:03d}"
        return f"{mushroom_name} - {tracking_number} - {restaurant_name} - Packed: {pack_date} - Shipped: {ship_date}"

    def clear_form(self):
        self.mushroom_type_var.set("")
        self.box_number_var.set("")
        self.restaurant_id_var.set("")
        self.pack_date_var.set("")
        self.ship_date_var.set("")

    def save_logs(self):
        try:
            with open(LOG_FILE, "w") as f:
                json.dump(self.logs, f, indent=4)
            self.show_toast("Logs saved successfully!", "success")
        except Exception as e:
            self.show_toast("Save failed", "error")

    def load_logs(self):
        if not os.path.exists("logs.json"):
            self.logs = []  # No file? Start fresh
            self.save_logs()
            self.update_filtered_logs()
            self.update_export_button_state()
            return

        try:
            with open("logs.json", "r") as f:
                loaded_data = json.load(f)

            if isinstance(loaded_data, list):
                self.logs = loaded_data
                self.show_toast("Logs loaded successfully!", "success")
            else:
                self.logs = []  # fallback to safe empty list
                self.show_toast("Invalid logs format detected! Logs reset.", "error")

            self.update_filtered_logs()
            self.update_export_button_state()

        except Exception as e:
            self.logs = []  # fallback to safe empty list
            self.show_toast(f"Failed to load logs: {e}", "error")
            self.update_filtered_logs()
            self.update_export_button_state()

    def clear_logs(self):
        if not self.logs:
            self.show_toast("No logs to clear.", "info")
            return

        confirm = messagebox.askyesno("Confirm", "Are you sure you want to delete all logs? A backup will be created.")
        if confirm:
            self.backup_logs()  # 🔥 Backup before clearing
            self.logs.clear()
            self.save_logs()
            self.update_filtered_logs()
            self.update_export_button_state()
            self.show_toast("All logs cleared successfully! Backup created.", "success")

    def backup_logs(self):
        base_folder = self.settings.get("export_folder", "")
        if not base_folder:
            base_folder = "."

        backup_folder = os.path.join(base_folder, "backups")

        # Create backups/ folder if it doesn't exist
        if not os.path.exists(backup_folder):
            os.makedirs(backup_folder)

        now = datetime.datetime.now().strftime("%Y-%m-%d_%H%M%S")
        backup_filename = os.path.join(backup_folder, f"logs_backup_{now}.json")

        try:
            with open(backup_filename, "w") as f:
                json.dump(self.logs, f, indent=4)
            self.show_toast(f"Backup created: {os.path.basename(backup_filename)}", "success")
        except Exception as e:
            self.show_toast(f"Failed to create backup: {e}", "error")

    def load_settings(self):
        if os.path.exists(self.settings_file):
            try:
                with open(self.settings_file, "r") as f:
                    self.settings.update(json.load(f))
                # Set theme immediately
                self.current_theme = self.settings.get("theme", "darkly")
                self.style.theme_use(self.current_theme)
            except Exception as e:
                self.show_toast(f"Failed to load settings: {e}", "error")

    def browse_logo_file(self):
        filepath = filedialog.askopenfilename(
            title="Select Logo Image",
            filetypes=[("Image Files", "*.png *.jpg *.jpeg")]
        )
        if filepath:
            self.logo_path_var.set(filepath)

    def save_settings(self):
        try:
            with open(self.settings_file, "w") as f:
                json.dump(self.settings, f, indent=4)
        except Exception as e:
            self.show_toast(f"Failed to save settings: {e}", "error")

    def save_settings_from_ui(self, window):
        self.settings["default_restaurant_id"] = self.default_restaurant_var.get()
        self.settings["invoice_template"] = self.invoice_template_var.get()
        self.settings["export_folder"] = self.export_folder_var.get()
        self.settings["default_export_format"] = self.default_export_format_var.get()
        self.settings["logo_path"] = self.logo_path_var.get()
        self.save_settings()
        self.show_toast("Settings saved!", "success")
        window.destroy()

    def open_settings_window(self):
        top = tk.Toplevel(self.root)
        top.title("Settings")
        top.geometry("450x350")
        top.resizable(False, False)

        # Default Restaurant ID
        ttk.Label(top, text="Default Restaurant ID:").pack(pady=(10, 0))
        self.default_restaurant_var = tk.IntVar(value=self.settings.get("default_restaurant_id", 1))
        ttk.Entry(top, textvariable=self.default_restaurant_var, width=10).pack()

        # Default Export Format
        ttk.Label(top, text="Default Export Format:").pack(pady=(10, 0))
        self.default_export_format_var = tk.StringVar(value=self.settings.get("default_export_format", "csv"))
        export_options = ["csv", "excel", "pdf"]
        ttk.Combobox(top, textvariable=self.default_export_format_var, values=export_options, state="readonly").pack()

        # Invoice Template Path
        ttk.Label(top, text="Invoice Template (.docx) Path:").pack(pady=(10, 0))
        self.invoice_template_var = tk.StringVar(value=self.settings.get("invoice_template", ""))

        file_frame = ttk.Frame(top)
        file_frame.pack(pady=(5, 0))
        ttk.Entry(file_frame, textvariable=self.invoice_template_var, width=30).pack(side="left", padx=(0, 5))
        ttk.Button(file_frame, text="Browse", command=self.browse_invoice_template).pack(side="left")

        # Logo Path
        ttk.Label(top, text="Logo Image Path (.png/.jpg):").pack(pady=(10, 0))
        self.logo_path_var = tk.StringVar(value=self.settings.get("logo_path", ""))

        logo_frame = ttk.Frame(top)
        logo_frame.pack(pady=(5, 0))
        ttk.Entry(logo_frame, textvariable=self.logo_path_var, width=30).pack(side="left", padx=(0, 5))
        ttk.Button(logo_frame, text="Browse Logo", command=self.browse_logo_file).pack(side="left")

        # Export Folder
        ttk.Label(top, text="Default Export Folder:").pack(pady=(10, 0))
        self.export_folder_var = tk.StringVar(value=self.settings.get("export_folder", ""))

        folder_frame = ttk.Frame(top)
        folder_frame.pack(pady=(5, 0))
        ttk.Entry(folder_frame, textvariable=self.export_folder_var, width=30).pack(side="left", padx=(0, 5))
        ttk.Button(folder_frame, text="Browse Folder", command=self.browse_export_folder).pack(side="left")

        # Save Button
        ttk.Button(top, text="Save Settings", command=lambda: self.save_settings_from_ui(top)).pack(pady=20)

    def update_filtered_logs(self, *_):
        search = self.search_var.get().lower()
        start_date = self.start_date_var.get()
        end_date = self.end_date_var.get()

        self.filtered_logs = []

        for log in self.logs:
            if search and search not in log.lower():
                continue

            parts = log.split(" - ")
            try:
                pack_date_str = parts[3].split(": ")[1]  # Extract Pack Date
            except (IndexError, ValueError):
                continue  # Skip bad logs

            if start_date:
                try:
                    if pack_date_str < start_date:
                        continue
                except:
                    continue

            if end_date:
                try:
                    if pack_date_str > end_date:
                        continue
                except:
                    continue

            self.filtered_logs.append(log)

    def view_log(self):
        if not self.filtered_logs:
            self.show_toast("No matching entries found.", "info")
        else:
            log_preview = "\n".join(self.filtered_logs[:5])
            self.show_toast(f"{len(self.filtered_logs)} result(s). Preview:\n{log_preview}", "info", duration=4000)

    def edit_logs(self):
        if not self.logs:
            self.show_toast("No logs to edit.", "info")
            return

        top = tk.Toplevel(self.root)
        top.title("Edit Logs")
        top.geometry("700x400")

        ttk.Label(top, text="Select a log to delete:").pack(pady=(10, 0))

        self.log_listbox = tk.Listbox(top, width=100, height=15)
        self.log_listbox.pack(pady=10)

        for log in self.logs:
            self.log_listbox.insert(tk.END, log)

        ttk.Button(top, text="Delete Selected", command=lambda: self.delete_selected_log(top)).pack(pady=10)

    def delete_selected_log(self, window):
        selection = self.log_listbox.curselection()
        if not selection:
            self.show_toast("No log selected.", "info")
            return

        confirm = messagebox.askyesno("Confirm", "Delete the selected log?")
        if not confirm:
            return

        index = selection[0]
        deleted_log = self.logs.pop(index)

        self.save_logs()
        self.update_filtered_logs()
        self.update_export_button_state()
        self.show_toast("Log deleted successfully.", "success")
        window.destroy()

    def show_charts(self):
        if not self.logs:
            messagebox.showwarning("No Data", "No entries to display charts.")
            return

        mushroom_counter = Counter()
        date_counter = Counter()

        for entry in self.logs:
            parts = entry.split(" - ")
            mushroom_type = parts[0]
            pack_date = parts[3].split(": ")[1]
            mushroom_counter[mushroom_type] += 1
            date_counter[pack_date] += 1

        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6))
        fig.suptitle('Delivery Statistics', fontsize=16)

        ax1.bar(mushroom_counter.keys(), mushroom_counter.values(), color='skyblue')
        ax1.set_title('Deliveries per Mushroom Type')
        ax1.set_xlabel('Mushroom Type')
        ax1.set_ylabel('Deliveries')
        ax1.tick_params(axis='x', rotation=45)

        sorted_dates = sorted(date_counter.items())
        dates, counts = zip(*sorted_dates)
        ax2.bar(dates, counts, color='lightgreen')
        ax2.set_title('Deliveries per Pack Date')
        ax2.set_xlabel('Pack Date')
        ax2.set_ylabel('Deliveries')
        ax2.tick_params(axis='x', rotation=90)

        plt.tight_layout()
        plt.show()

    def show_toast(self, message, type="info", duration=3000):
        toast = tk.Toplevel(self.root)
        toast.overrideredirect(True)
        toast.attributes("-topmost", True)

        colors = {
            "info": "#2196F3",      # blue
            "success": "#4CAF50",   # green
            "error": "#F44336"      # red
        }

        bg = colors.get(type, "#2196F3")

        label = tk.Label(toast, text=message, bg=bg, fg="white",
                        font=("Segoe UI", 10), padx=20, pady=10)
        label.pack()

        # Position: bottom-right of main window
        self.root.update_idletasks()
        x = self.root.winfo_x() + self.root.winfo_width() - label.winfo_reqwidth() - 40
        y = self.root.winfo_y() + self.root.winfo_height() - label.winfo_reqheight() - 60
        toast.geometry(f"+{x}+{y}")

        # Auto-destroy
        toast.after(duration, toast.destroy)

    def export_to_csv(self):
        if not self.logs:
            self.show_toast("No data to export.", "error")
            return

        folder = self.settings.get("export_folder", "")
        if not folder:
            folder = "."

        today = datetime.date.today().strftime("%Y-%m-%d")
        filepath = os.path.join(folder, f"traceability_log_{today}.csv")

        with open(filepath, "w", newline="") as file:
            writer = csv.writer(file)
            writer.writerow(["Mushroom Type", "Box Number", "Restaurant Name", "Packed Date", "Shipped Date"])
            for entry in self.logs:
                parts = entry.split(" - ")
                writer.writerow([
                    parts[0],
                    parts[1].split("BOX")[1],
                    parts[2],
                    parts[3].split(": ")[1],
                    parts[4].split(": ")[1]
                ])
        self.show_toast(f"Saved to {filepath}!", "success")

    def export_to_excel(self):
        if not self.logs:
            self.show_toast("No data to export.", "error")
            return

        folder = self.settings.get("export_folder", "")
        if not folder:
            folder = "."

        today = datetime.date.today().strftime("%Y-%m-%d")
        filepath = os.path.join(folder, f"traceability_log_{today}.xlsx")

        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Traceability Log"
        sheet.append(["Mushroom Type", "Box Number", "Restaurant Name", "Packed Date", "Shipped Date"])

        for entry in self.logs:
            parts = entry.split(" - ")
            sheet.append([
                parts[0],
                parts[1].split("BOX")[1],
                parts[2],
                parts[3].split(": ")[1],
                parts[4].split(": ")[1]
            ])

        workbook.save(filepath)
        self.show_toast(f"Saved to {filepath}!", "success")

    def export_summary_report(self):
        if not self.logs:
            self.show_toast("No data to export summary.", "error")
            return

        folder = self.settings.get("export_folder", "")
        if not folder:
            folder = "."

        today = datetime.date.today().strftime("%Y-%m-%d")
        doc_name = os.path.join(folder, f"summary_report_{today}.docx")
        pdf_name = os.path.join(folder, f"summary_report_{today}.pdf")

        from collections import Counter
        import docx.shared

        # Count mushrooms
        mushroom_counter = Counter()
        for entry in self.logs:
            parts = entry.split(" - ")
            mushroom_type = parts[0]
            mushroom_counter[mushroom_type] += 1

        doc = Document()

        # Insert logo first if available
        try:
            logo_path = self.settings.get("logo_path", "")
            if logo_path and os.path.exists(logo_path):
                absolute_logo_path = os.path.abspath(logo_path)
                from PIL import Image

                # Open image with PIL to check size
                img = Image.open(absolute_logo_path)
                img_width, img_height = img.size  # in pixels

                # Define a maximum width (in inches)
                max_width_inch = 2.0
                dpi = img.info.get('dpi', (96, 96))[0]  # default 96 dpi if missing

                # Convert pixels to inches
                img_width_inch = img_width / dpi

                # If image is wider than allowed, resize it
                if img_width_inch > max_width_inch:
                    doc.add_picture(absolute_logo_path, width=docx.shared.Inches(max_width_inch))
                else:
                    doc.add_picture(absolute_logo_path)  # Insert without resizing
        except Exception as e:
            self.show_toast(f"Failed to insert logo: {e}", "error")

        # Add Title and Metadata
        doc.add_heading('Mushroom Deliveries Summary', 0)
        doc.add_paragraph(f"Export Date: {today}")
        doc.add_paragraph(f"Total Deliveries: {len(self.logs)}")

        # Deliveries per Mushroom Type
        doc.add_heading('Deliveries per Mushroom Type', level=1)
        for mushroom, count in mushroom_counter.items():
            doc.add_paragraph(f"{mushroom}: {count} deliveries", style="List Bullet")

        # Detailed Deliveries Table
        doc.add_heading('Detailed Deliveries', level=1)

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        headers = ["Mushroom Type", "Box Number", "Restaurant", "Pack Date", "Ship Date"]
        for i, header in enumerate(headers):
            hdr[i].text = header

        for entry in self.logs:
            parts = entry.split(" - ")
            row = table.add_row().cells
            row[0].text = parts[0]
            row[1].text = parts[1].split("BOX")[1]
            row[2].text = parts[2]
            row[3].text = parts[3].split(": ")[1]
            row[4].text = parts[4].split(": ")[1]

        doc.save(doc_name)

        # Convert to PDF
        try:
            from docx2pdf import convert
            convert(doc_name, pdf_name)
            self.show_toast(f"Summary exported: {os.path.basename(pdf_name)}", "success")
            os.startfile(pdf_name)
        except Exception as e:
            self.show_toast(f"PDF conversion failed: {e}", "error")

    def update_export_button_state(self):
        if hasattr(self, 'export_button'):
            if self.logs:
                self.export_button.config(state="normal")
            else:
                self.export_button.config(state="disabled")

    def generate_invoice(self):
        if not self.logs:
            self.show_toast("No data to generate invoice.", "error")
            return

        folder = self.settings.get("export_folder", "")
        if not folder:
            folder = "."

        today = datetime.date.today().strftime("%Y-%m-%d")
        doc_name = os.path.join(folder, f"invoice_{today}.docx")
        pdf_name = os.path.join(folder, f"invoice_{today}.pdf")

        doc = Document()
        doc.add_heading('Mushroom Traceability Invoice', 0)

        # Table header
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        headers = ["Mushroom Type", "Box Number", "Restaurant", "Pack Date", "Ship Date"]
        for i, header in enumerate(headers):
            hdr[i].text = header

        # Only include most recent entry
        latest_entry = self.logs[-1] if self.logs else None
        if latest_entry:
            parts = latest_entry.split(" - ")
            row = table.add_row().cells
            row[0].text = parts[0]
            row[1].text = parts[1].split("BOX")[1]
            row[2].text = parts[2]
            row[3].text = parts[3].split(": ")[1]
            row[4].text = parts[4].split(": ")[1]

        doc.save(doc_name)

        try:
            from docx2pdf import convert
            convert(doc_name, pdf_name)
            self.show_toast(f"Invoice saved as {os.path.basename(pdf_name)}", "success")
            os.startfile(pdf_name)
        except Exception as e:
            self.show_toast(f"PDF conversion failed: {e}", "error")

    def toggle_mode(self):
        new_value = "0" if self.is_mock_mode else "1"
        os.environ["USE_MOCK_SQUARE"] = new_value
        if messagebox.askyesno("Restart Required", "Restart app to apply mode switch?"):
            python = sys.executable
            os.execl(python, python, *sys.argv)

    def get_mode_text(self):
        return f"🧪 Mock Mode: {'ON' if self.is_mock_mode else 'OFF'}"

    def browse_invoice_template(self):
        filepath = filedialog.askopenfilename(
            title="Select Invoice Template",
            filetypes=[("Word Documents", "*.docx")]
        )
        if filepath:
            self.invoice_template_var.set(filepath)

    def browse_export_folder(self):
        folder_path = filedialog.askdirectory(
            title="Select Export Folder"
        )
        if folder_path:
            self.export_folder_var.set(folder_path)

    def export_data(self):
        if not self.logs:
            self.show_toast("No data to export.", "error")
            return

        preferred_format = self.settings.get("default_export_format", "csv")
        if preferred_format == "csv":
            self.export_to_csv()
        elif preferred_format == "excel":
            self.export_to_excel()
        elif preferred_format == "pdf":
            self.generate_invoice()
        else:
            self.show_toast(f"Unknown export format: {preferred_format}", "error")

    def restore_backup(self):
        base_folder = self.settings.get("export_folder", "")
        if not base_folder:
            base_folder = "."

        backup_folder = os.path.join(base_folder, "backups")
        if not os.path.exists(backup_folder):
            os.makedirs(backup_folder)

        file_path = filedialog.askopenfilename(
            initialdir=backup_folder,
            title="Select Backup File",
            filetypes=[("JSON Backup Files", "*.json")]
        )
        if not file_path:
            return  # User canceled

        try:
            with open(file_path, "r") as f:
                restored_data = json.load(f)

            if isinstance(restored_data, list):
                self.logs = restored_data
                self.save_logs()
                self.update_filtered_logs()
                self.update_export_button_state()
                self.show_toast(f"Backup restored successfully from {os.path.basename(file_path)}!", "success")
            else:
                self.logs = []  # Clear to safe empty list
                self.show_toast("Invalid backup file format! Logs reset.", "error")

        except Exception as e:
            self.show_toast(f"Failed to restore backup: {e}", "error")

    def open_backup_manager(self):
        backup_window = tk.Toplevel(self.root)
        backup_window.title("Backup Manager")
        backup_window.geometry("500x400")
        backup_window.resizable(False, False)

        ttk.Label(backup_window, text="Available Backups:").pack(pady=(10, 0))

        self.backup_listbox = tk.Listbox(backup_window, width=60, height=15)
        self.backup_listbox.pack(pady=10)

        button_frame = ttk.Frame(backup_window)
        button_frame.pack(pady=10)

        ttk.Button(button_frame, text="Restore Selected", command=self.restore_selected_backup).pack(side="left", padx=10)
        ttk.Button(button_frame, text="Delete Selected", command=self.delete_selected_backup).pack(side="left", padx=10)

        # Load available backups into listbox
        self.load_backup_list()

    def load_backup_list(self):
        self.backup_listbox.delete(0, tk.END)
        folder = os.path.join(self.settings.get("export_folder", "."), "backups")
        if not os.path.exists(folder):
            os.makedirs(folder)

        files = [f for f in os.listdir(folder) if f.endswith(".json")]
        for f in sorted(files, reverse=True):
            self.backup_listbox.insert(tk.END, f)

    def restore_selected_backup(self):
        selection = self.backup_listbox.curselection()
        if not selection:
            self.show_toast("No backup selected.", "info")
            return

        selected_file = self.backup_listbox.get(selection[0])
        folder = os.path.join(self.settings.get("export_folder", "."), "backups")
        filepath = os.path.join(folder, selected_file)

        try:
            with open(filepath, "r") as f:
                restored_logs = json.load(f)

            if isinstance(restored_logs, list):
                self.logs = restored_logs
                self.save_logs()
                self.update_filtered_logs()
                self.update_export_button_state()
                self.show_toast(f"Backup restored from {selected_file}!", "success")
            else:
                self.show_toast("Invalid backup file format!", "error")

        except Exception as e:
            self.show_toast(f"Restore failed: {e}", "error")

if __name__ == "__main__":
    root = tk.Tk()
    app = MushroomApp(root)
    root.mainloop()
