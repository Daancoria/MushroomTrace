from docx import Document
import datetime
from config import (
    MUSHROOM_TYPES,
    RESTAURANT_ASSIGNMENTS,
    SQUARE_ACCESS_TOKEN,
    SQUARE_LOCATION_ID,
    SQUARE_ORDER_ID,
    SQUARE_CUSTOMER_ID,
    USE_MOCK_SQUARE,
)

# Attempt to import real Square client
try:
    from square.client import Client as RealClient
except ImportError:
    RealClient = None

# --- Mock Square Client ---
class MockSquareClient:
    def __init__(self, access_token=None):
        self.invoices = self.MockInvoices()

    class MockInvoices:
        def create_invoice(self, body):
            print(f"[MOCK] Square invoice created for: {body['invoice']['description']}")
            return self.MockResponse()

        class MockResponse:
            def is_success(self):
                return True

            @property
            def errors(self):
                return None

# Choose client based on toggle
Client = MockSquareClient if USE_MOCK_SQUARE else RealClient


class TraceabilityManager:
    def __init__(self):
        self.logs = []
        self.client = Client(access_token=SQUARE_ACCESS_TOKEN)

    def generate_tracking_label(self, mushroom_type, box_number, restaurant_id, pack_date, ship_date):
        mushroom_name = MUSHROOM_TYPES.get(mushroom_type)
        if not mushroom_name:
            raise ValueError("Invalid mushroom type selected.")
        restaurant_name = RESTAURANT_ASSIGNMENTS.get(restaurant_id)
        if not restaurant_name:
            raise ValueError("Invalid restaurant ID.")
        
        tracking_number = f"{pack_date}-BOX{box_number:03d}"
        label = f"{mushroom_name} - {tracking_number} - {restaurant_name} - Packed: {pack_date} - Shipped: {ship_date}"
        self.logs.append(label)
        return label

    def generate_invoice_doc(self, filename="invoice.docx"):
        current_date = datetime.datetime.now().strftime("%Y-%m-%d")
        doc = Document()
        doc.add_heading('Mushroom Traceability Label and Invoice', 0)
        doc.add_heading('Traceability Labels', level=1)

        for label in self.logs:
            doc.add_paragraph(label)

        doc.add_heading('Invoice Details', level=1)
        table = doc.add_table(rows=len(self.logs)+1, cols=5)
        table.style = 'Table Grid'

        headers = ['Mushroom Type', 'Box Number', 'Restaurant Name', 'Pack Date', 'Ship Date']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        for i, label in enumerate(self.logs, start=1):
            parts = label.split(' - ')
            mushroom_type = parts[0]
            box_number = parts[1].split('BOX')[1]
            restaurant_name = parts[2]
            pack_date = parts[3].split(': ')[1]
            ship_date = parts[4].split(': ')[1]

            row = table.rows[i].cells
            row[0].text = mushroom_type
            row[1].text = box_number
            row[2].text = restaurant_name
            row[3].text = pack_date
            row[4].text = ship_date

        doc.save(filename)

    def create_square_invoices(self):
        current_date = datetime.datetime.now().strftime("%Y-%m-%d")
        for label in self.logs:
            parts = label.split(' - ')
            box_number = parts[1].split('BOX')[1]
            restaurant_name = parts[2]

            invoice_data = {
                "idempotency_key": f"{current_date}-{box_number}",
                "invoice": {
                    "location_id": SQUARE_LOCATION_ID,
                    "order_id": SQUARE_ORDER_ID,
                    "primary_recipient": {
                        "customer_id": SQUARE_CUSTOMER_ID
                    },
                    "payment_requests": [{
                        "request_type": "BALANCE",
                        "due_date": current_date,
                        "fixed_amount_requested_money": {"amount": 0, "currency": "USD"}
                    }],
                    "delivery_method": "EMAIL",
                    "invoice_number": f"INV-{current_date}-{box_number}",
                    "title": "Mushroom Invoice",
                    "description": f"Invoice for {label} delivered to {restaurant_name}"
                }
            }

            result = self.client.invoices.create_invoice(body=invoice_data)

            if not result.is_success():
                raise Exception(f"[Invoice Error] {result.errors}")
