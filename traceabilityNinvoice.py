import datetime
import csv
import tkinter as tk
from tkinter import ttk, messagebox
import os
import subprocess
from docx import Document
from square.client import Client

# Sample restaurant assignments
restaurant_assignments = {
    1: "Restaurant A",
    2: "Restaurant B",
    3: "Restaurant C"
}

# Square API credentials from environment variables
square_access_token = os.getenv('SQUARE_ACCESS_TOKEN')
location_id = os.getenv('SQUARE_LOCATION_ID')
order_id = os.getenv('SQUARE_ORDER_ID')
customer_id = os.getenv('SQUARE_CUSTOMER_ID')

# Initialize Square client
client = Client(access_token=square_access_token)

# List to store traceability logs
traceability_logs = []

def generate_tracking_number(mushroom_type, box_number, restaurant_name, pack_date, ship_date):
    if mushroom_type == 1:
        mushroom_name = "Blue Oyster"
    elif mushroom_type == 2:
        mushroom_name = "Lion's Mane"
    else:
        return "Invalid mushroom type"

    tracking_number = f"{pack_date}-BOX{box_number:03d}"
    traceability_label = f"{mushroom_name} - {tracking_number} - {restaurant_name} - Packed: {pack_date} - Shipped: {ship_date}"

    return traceability_label

def add_to_traceability_log():
    try:
        mushroom_type = int(mushroom_type_var.get())
        box_number = int(box_number_var.get())
        restaurant_id = int(restaurant_id_var.get())
        pack_date = pack_date_var.get()
        ship_date = ship_date_var.get()

        if mushroom_type not in [1, 2]:
            messagebox.showerror("Input Error", "Invalid mushroom type. Please enter 1 for Blue Oyster or 2 for Lion's Mane.")
            return

        if restaurant_id not in restaurant_assignments:
            messagebox.showerror("Input Error", "Invalid restaurant ID.")
            return

        restaurant_name = restaurant_assignments[restaurant_id]
        label = generate_tracking_number(mushroom_type, box_number, restaurant_name, pack_date, ship_date)
        messagebox.showinfo("Traceability Label", f"Traceability Label: {label}\nRestaurant: {restaurant_name}")
        
        # Add to traceability logs
        traceability_logs.append(label)
        messagebox.showinfo("Success", "Entry added to traceability log.")
    except ValueError:
        messagebox.showerror("Input Error", "Please enter valid numbers for mushroom type, box number, and restaurant ID.")
    except Exception as e:
        messagebox.showerror("Error", f"An unexpected error occurred: {e}")

def show_traceability_log():
    if not traceability_logs:
        messagebox.showinfo("Traceability Log", "The traceability log is currently empty.")
    else:
        log_text = "\n".join(traceability_logs)
        messagebox.showinfo("Traceability Log", f"Current Traceability Log:\n{log_text}")

def generate_and_print_invoice():
    try:
        if not traceability_logs:
            messagebox.showerror("Error", "No entries in traceability log.")
            return
        
        # Generate the invoice and save to Word document
        generate_invoice()
        
        # Open the Word document for editing
        open_with_word("invoice.docx")
        
        # Wait for user to press Enter to print
        input("Press Enter to print the invoice...")
        
        # Print the Word document
        print_document("invoice.docx")
    except Exception as e:
        messagebox.showerror("Error", f"An unexpected error occurred: {e}")

def generate_invoice():
    current_date = datetime.datetime.now().strftime("%Y-%m-%d")
    produce_cost = ""  # Leave space for produce cost to be filled in later

    # Create a Word document
    doc = Document()
    doc.add_heading('Mushroom Traceability Label and Invoice', 0)

    doc.add_heading('Traceability Label', level=1)
    for label in traceability_logs:
        doc.add_paragraph(label)

    doc.add_heading('Invoice', level=1)
    table = doc.add_table(rows=len(traceability_logs) + 1, cols=8)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Invoice Date'
    hdr_cells[1].text = 'Box Number'
    hdr_cells[2].text = 'Mushroom Type'
    hdr_cells[3].text = 'Traceability Label'
    hdr_cells[4].text = 'Produce Cost'
    hdr_cells[5].text = 'Restaurant ID'
    hdr_cells[6].text = 'Restaurant Name'
    hdr_cells[7].text = 'Pack Date'
    hdr_cells[8].text = 'Ship Date'

    for i, label in enumerate(traceability_logs, start=1):
        parts = label.split(' - ')
        mushroom_type = parts[0]
        box_number = parts[1].split('BOX')[1]
        restaurant_name = parts[2]
        pack_date = parts[3].split(': ')[1]
        ship_date = parts[4].split(': ')[1]
        restaurant_id = [key for key, value in restaurant_assignments.items() if value == restaurant_name][0]

        row_cells = table.rows[i].cells
        row_cells[0].text = current_date
        row_cells[1].text = box_number
        row_cells[2].text = mushroom_type
        row_cells[3].text = label
        row_cells[4].text = produce_cost
        row_cells[5].text = str(restaurant_id)
        row_cells[6].text = restaurant_name
        row_cells[7].text = pack_date
        row_cells[8].text = ship_date

    doc.save('invoice.docx')
    messagebox.showinfo("Invoice Generated", "Invoice has been generated and saved as 'invoice.docx'.")

    # Create an invoice using Square API
    create_square_invoice()

def create_square_invoice():
    current_date = datetime.datetime.now().strftime("%Y-%m-%d")

    for label in traceability_logs:
        parts = label.split(' - ')
        box_number = parts[1].split('BOX')[1]
        mushroom_type = parts[0]
        restaurant_name = parts[2]

        # Create invoice data
        invoice_data = {
            "idempotency_key": f"{current_date}-{box_number}",
            "invoice": {
                "location_id": location_id,
                "order_id": order_id,
                "primary_recipient": {
                    "customer_id": customer_id
                },
                "payment_requests": [
                    {
                        "request_type": "BALANCE",
                        "due_date": current_date,
                        "fixed_amount_requested_money": {
                            "amount": 0,  # Amount in cents, update as needed
                            "currency": "USD"
                        }
                    }
                ],
                "delivery_method": "EMAIL",
                "invoice_number": f"INV-{current_date}-{box_number}",
                "title": "Mushroom Invoice",
                "description": f"Invoice for {label} delivered to {restaurant_name}"
            }
        }

        # Call Square API to create the invoice
        result = client.invoices.create_invoice(body=invoice_data)

        if result.is_success():
            messagebox.showinfo("Square Invoice", f"Square invoice for {label} has been created successfully.")
        else:
            messagebox.showerror("Square Invoice Error", f"Failed to create Square invoice for {label}: {result.errors}")

def print_document(file_path):
    os.startfile(file_path, "print")

def print_traceability_logs():
    # Save the traceability logs to a temporary text file
    with open("traceability_logs.txt", "w") as file:
        for log in traceability_logs:
            file.write(log + "\n")
    
    # Open the text file in the default text editor (e.g., Notepad)
    os.startfile("traceability_logs.txt")

def open_with_word(file_path):
    # Open the file with Microsoft Word explicitly
    subprocess.run(['start', 'winword', file_path], shell=True)

# Create the GUI window
root = tk.Tk()
root.title("Mushroom Tracking Label and Invoice Generator")

# Mushroom Type
ttk.Label(root, text="Mushroom Type (1 for Blue Oyster, 2 for Lion's Mane):").grid(column=0, row=0, padx=10, pady=5)
mushroom_type_var = tk.StringVar()
ttk.Entry(root, textvariable=mushroom_type_var).grid(column=1, row=0, padx=10, pady=5)

# Box Number
ttk.Label(root, text="Box Number:").grid(column=0, row=1, padx=10, pady=5)
box_number_var = tk.StringVar()
ttk.Entry(root, textvariable=box_number_var).grid(column=1, row=1, padx=10, pady=5)

# Restaurant ID
ttk.Label(root, text="Restaurant ID:").grid(column=0, row=2, padx=10, pady=5)
restaurant_id_var = tk.StringVar()
ttk.Entry(root, textvariable=restaurant_id_var).grid(column=1, row=2, padx=10, pady=5)

# Pack Date
ttk.Label(root, text="Pack Date (YYYY-MM-DD):").grid(column=0, row=3, padx=10, pady=5)
pack_date_var = tk.StringVar()
ttk.Entry(root, textvariable=pack_date_var).grid(column=1, row=3, padx=10, pady=5)

# Ship Date
ttk.Label(root, text="Ship Date (YYYY-MM-DD):").grid(column=0, row=4, padx=10, pady=5)
ship_date_var = tk.StringVar()
ttk.Entry(root, textvariable=ship_date_var).grid(column=1, row=4, padx=10, pady=5)

# Add to Traceability Log Button
ttk.Button(root, text="Add to Traceability Log", command=add_to_traceability_log).grid(column=0, row=5, columnspan=2, padx=10, pady=10)

# Show Traceability Log Button
ttk.Button(root, text="Show Traceability Log", command=show_traceability_log).grid(column=0, row=6, columnspan=2, padx=10, pady=10)

# Generate and Print Invoice Button
ttk.Button(root, text="Generate and Print Invoice", command=generate_and_print_invoice).grid(column=0, row=7, columnspan=2, padx=10, pady=10)

# Print Traceability Logs Button
ttk.Button(root, text="Print Traceability Logs", command=print_traceability_logs).grid(column=0, row=8, columnspan=2, padx=10, pady=10)

root.mainloop()
